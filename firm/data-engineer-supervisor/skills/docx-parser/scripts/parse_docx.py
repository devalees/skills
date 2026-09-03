import os
import sys
import json
import docx

def parse_docx(file_path):
    doc = docx.Document(file_path)
    paragraphs = []
    tables_data = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            paragraphs.append(txt)

    for t_idx, table in enumerate(doc.tables):
        raw_rows = []
        for row in table.rows:
            raw_rows.append([cell.text.strip() for cell in row.cells])
        
        if not raw_rows:
            continue
        
        headers = raw_rows[0]
        data_rows = []
        for r in raw_rows[1:]:
            row_dict = {}
            for col_idx, h in enumerate(headers):
                col_key = h if h != "" else f"col_{col_idx+1}"
                row_dict[col_key] = r[col_idx] if col_idx < len(r) else ""
            data_rows.append(row_dict)

        tables_data.append({
            "table_index": t_idx + 1,
            "headers": headers,
            "row_count": len(data_rows),
            "rows": data_rows
        })

    structure = "hybrid" if tables_data and paragraphs else ("tabular" if tables_data else "narrative_text")

    envelope = {
        "document_metadata": {
            "file_name": os.path.basename(file_path),
            "file_type": "docx",
            "structure_type": structure,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables_data)
        },
        "payload": {
            "paragraphs": paragraphs,
            "tables": tables_data
        },
        "data_health": {
            "total_paragraphs": len(paragraphs),
            "total_tables": len(tables_data),
            "status": "success"
        }
    }
    return envelope

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 parse_docx.py <path_to_docx_file>")
        sys.exit(1)
    result = parse_docx(sys.argv[1])
    print(json.dumps(result, ensure_ascii=False, indent=2))
